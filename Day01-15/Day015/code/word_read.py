"""
读取Word文件

@Author:jyang
@Date:5/22/2019
"""

from docx import Document

doc = Document('../res/E12Cons.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[-1].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for p in doc.paragraphs:
    content.append(p.text)
print(''.join(content))